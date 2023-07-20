# #엑셀에 붙여넣기
# pip install selenium
# pip install pyautogui
# pip install python-docx
# pip install pywin32
# pip install pywinauto
# pip install webdriver-manager
# pip install lxml
# pip install html_parser
# pip install bs4
# pip install pandas

import time
from issues_in_stock import yesm,name_fi, removeff
from docx import Document
import win32com.client, os
import pandas as pd
import requests
from bs4 import BeautifulSoup as bs4
from urllib.request import Request, urlopen
import openpyxl

this_program_directory = os.path.dirname(os.path.abspath(__file__))
os.chdir(this_program_directory)

document_path = './DB\\' + name_fi.split('\\')[4].replace("hwp", "docx")
# document_path = './DB\\' + '7.19특징주.docx' #for test
# print(document_path)
doc = Document(document_path)    #day_docx


table_first = doc.tables[0] #특징주
table_second = doc.tables[1] #시간외
name_up=[]

impact = []
impact_i=0
for r in table_first.rows:
    data_list_first = []
    for cell in r.cells:
        for para in cell.paragraphs:
            # print(para.text)
            data_list_first.append(para.text)
    print(', '.join(data_list_first))
    if impact_i !=0:
        impact.append(data_list_first)
        name_up.append(data_list_first[0])

    else:
        impact_i+=1



#비어있는 데이터프레임 생성
stock_name = pd.DataFrame()
temp_name = pd.DataFrame()

#장중 특징주 종목명 모음
for i in range(0,len(impact)):

    temp_ = impact[i][0]
    temp_name = pd.DataFrame({'종목명': [temp_]})
    stock_name = pd.concat([stock_name, temp_name], axis=0)

impact_result = stock_name.reset_index(drop=True)




for i in impact:
    i.insert(0, '정규장')
    # for j in i:
        # print(j, end="|")
    # print()

print()
print("----------------")
print()


over_t = []
over_t_i=0
for row in table_second.rows:
    data_list_second = []
    for cell in row.cells:
        #docx에서 한 글자씩 떼기 위함
        for para in cell.paragraphs:
            # print(para.text)
            data_list_second.append(para.text)
    print(', '.join(data_list_second))
    if over_t_i !=0:
        over_t.append(data_list_second)
        name_up.append(data_list_first[0])
    else:
        over_t_i+=1

#비어있는 데이터프레임 생성
stock_name = pd.DataFrame()
temp_name = pd.DataFrame()

# 장중 특징주 종목명 모음
for i in range(0, len(over_t)):
    temp_ = over_t[i][0]
    temp_name = pd.DataFrame({'종목명': [temp_]})
    stock_name = pd.concat([stock_name, temp_name], axis=0)

#시간외 단일가 특징주 종목들
overt_result = stock_name.reset_index(drop=True)

#종목명들을 모은 결과물
index_result = pd.concat([impact_result, overt_result], axis=0)
index_result = index_result.reset_index(drop=True)

# print(index_result)


for i in over_t:
    i.insert(0, '시간외')
    # for j in i:
    #     print(j, end="|")
    # print()

# 데이터를 잘 긁어오는것을 확인완료

# 종목별 테마 인덱스 크롤링

# 결과물이 들어갈 판다스
result = pd.DataFrame()
result_label = pd.DataFrame()

for i in range(0, len(index_result)):

    #여기에 테마 내용들이 쭉 적히게 될 것
    index = ''

    # 종목명
    keyword = index_result['종목명'][i]
    print(keyword)

    try:
    # 크롤링 및 파싱
        url = f'https://finance.finup.co.kr/Stock/{keyword}'
        req = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'})
        parsing = bs4(req.text, 'html.parser')

        # 파싱한것들중에 테마만 긁기
        label_ = parsing.find_all(class_='label')

        #테마 이름들을 한 줄로 합치기
        for i in range(6, len(label_)):
            index = index + label_[i].text + ', '

        #테마 이름들을 temp_pd라는 데이타프레임에 입력
        temp_pd = pd.DataFrame({'종목명': [keyword], '테마': [index[:-2]]})
        result = pd.concat([result, temp_pd], axis=0)


    except:
        pass

print('ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ')
print(result)
print('ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ')

#엑셀파일 경로
excel = win32com.client.Dispatch("Excel.Application")
excel.Visible = True
print("os.getcwd", os.getcwd())
excel_file = this_program_directory+'\\DB' + '\\특징주DB.xlsm'
print("excel_path : ", excel_file)

result = result.reset_index(drop=True)
result = result.drop_duplicates(['종목명'], keep='last', ignore_index=True)

#테마 업데이트
theme_update = pd.DataFrame()
#원래 있던 엑셀파일에서 theme 시트만 불러오기
theme_data = pd.read_excel(excel_file, sheet_name='theme')
#concat 시키고 중복값 제거
theme_update = pd.concat([theme_data,result], axis=0)
theme_update = theme_update.drop_duplicates(['종목명'], keep='last', ignore_index=True)


#판다스 데이타프레임을 리스트형식으로 변환
theme_list = theme_update.values.tolist()
#theme_list는 첫 행에 종목명, 테마가 없기 때문에 이를 삽입

print(theme_list)







excel = win32com.client.Dispatch("Excel.Application")
excel.Visible = True
print("os.getcwd", os.getcwd())
excel_file = this_program_directory+'\\DB' + '\\특징주DB.xlsm'
print("excel_path : ", excel_file)
# wb = excel.Workbooks.Add()  #엑셀 프로그램에 Workbook 추가(객체 생성)
os.chdir(this_program_directory)
wb = excel.Workbooks.Open(excel_file)


index0 = str(yesm.year) + str(yesm.month).zfill(2) + str(yesm.day).zfill(2)
def copycells(ia, ib, array):
    a = ia
    for i in array:
        b = ib
        for j in i:
            ws_DB.cells(a, b).Value = str(j)
            ws_DB.cells(a, ib-2).Value = index0
            b += 1
        a += 1
    return a

def copycellss(ia, ib, array):
    a = ia
    for i in array:
        b = ib
        for j in i:
            if str([j]).strip("[""]""\'") != '종목명':
                ws_them.cells(a,b).Value = str([j]).strip("[""]""\'")
            b += 1
        a += 1
    return a

#ws_temp.Delete()


ws_DB = wb.Worksheets("특징주DB")  #-시트 이름으로 객체 설정
ws_them = wb.Worksheets("theme")
ws_DB.Select()
time.sleep(0.1)

data_range = ws_DB.UsedRange()
alpha=1
beta=1
for i in data_range:
    row_data = []
    for j in i:
        cell_value = j
        row_data.append(str(cell_value))
        # print(row_data[0])
    # print(', '.join(row_data))
    alpha +=1
    beta = row_data[0]
    if(row_data[1] < '0'):
        break

print(alpha)
first_tabel=copycells(alpha-1, 3, impact)
copycells(first_tabel,3,over_t)
copycellss(2,1, theme_list)
save_path = this_program_directory + '\\DB'
removeff(document_path)

#theme_upload
ws_them.Select()

wb.SaveAs(save_path + "/특징주DB.xlsm")


