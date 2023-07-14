from docx import Document

document = Document()

document.add_heading('코딩유치원 python-docx 강의', level = 0)

p = document.add_paragraph('안녕하세요, 코린이 여러분!')
p.add_run(' 코딩유치원에 오신 것을 환영합니다.').bold = True

document.add_paragraph('문장 추가 1')
document.add_paragraph('문장 추가 2')
document.add_paragraph('문장 추가 3')
document.add_paragraph('문장 추가 4')

records = (
    (1, '하나', 'one'),
    (2, '둘', 'two'),
    (3, '셋', 'three')
)

table = document.add_table(rows=1, cols=3)

# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid'로 설정
table.style = document.styles['Table Grid']

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = '한국어'
hdr_cells[2].text = '영어'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.save('예제 문서.docx')