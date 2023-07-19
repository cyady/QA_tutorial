#엑셀에 붙여넣기
#pip install python-docx
#pip install pywin32
#pip install pywinauto
#pip install webdriver-manager
#pip install lxml
#pip install html_parser

import time
from issues_in_stock import yesm,name_fi, removeff
from docx import Document
import win32com.client, os

this_program_directory = os.path.dirname(os.path.abspath(__file__))
os.chdir(this_program_directory)

document_path = './DB\\' + name_fi.split('\\')[4].replace("hwp", "docx")
print(document_path)
doc = Document(document_path)    #day_docx


table_first = doc.tables[0] #특징주
table_second = doc.tables[1] #시간외

impact = []
impact_i=0
for r in table_first.rows:
    data_list_first = []
    for cell in r.cells:
        for para in cell.paragraphs:
            # print(para.text)
            data_list_first.append(para.text)
    if impact_i !=0:
        impact.append(data_list_first)
    else:
        impact_i+=1

for i in impact:
    i.insert(0, '정규장')
    for j in i:
        print(j, end="|")
    print()
print()
print("----------------")
print()


over_t = []
for row in table_second.rows:
    data_list_second = []
    for cell in row.cells:
        for para in cell.paragraphs:
            # print(para.text)
            data_list_second.append(para.text)
    # print(', '.join(data_list_second))
    over_t.append(data_list_second)

for i in over_t:
    i.insert(0, '시간외')
    for j in i:
        print(j, end="|")
    print()

# 데이터를 잘 긁어오는것을 확인완료



excel = win32com.client.Dispatch("Excel.Application")
excel.Visible = True
print("os.getcwd", os.getcwd())
excel_file = this_program_directory+'\\DB' + '\\특징주DB.xlsm'
print("excel_path : ", excel_file)
# wb = excel.Workbooks.Add()  #엑셀 프로그램에 Workbook 추가(객체 생성)
os.chdir(this_program_directory)
wb = excel.Workbooks.Open(excel_file)


index0 = str(yesm.year) + str(yesm.month).zfill(2) + str(yesm.day).zfill(2)
def copycells(ia, ib, array):
    a = ia
    for i in array:
        b = ib
        for j in i:
            ws_DB.cells(a, b).Value = str(j)
            ws_DB.cells(a, ib-2).Value = index0
            b += 1
        a += 1
    return a

##
# ws_temp.Delete()


ws_DB = wb.Worksheets("특징주DB")  #-시트 이름으로 객체 설정
ws_DB.Select()
time.sleep(0.1)

data_range = ws_DB.UsedRange()
alpha=1
beta=1
for i in data_range:

    row_data = []
    for j in i:
        cell_value = j
        row_data.append(str(cell_value))
        # print(row_data[0])
    print(', '.join(row_data))
    alpha +=1
    beta = row_data[0]
    if(row_data[1] < '0'):
        break
print(alpha)
first_tabel=copycells(alpha-1, 3, impact)
copycells(first_tabel,3,over_t)
save_path = this_program_directory + '\\DB'
removeff(document_path)

wb.SaveAs(save_path + "/특징주DB.xlsm")


